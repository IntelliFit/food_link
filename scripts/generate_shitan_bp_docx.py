from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "食探（智健食探）商业计划书.docx"

FONT_CN = "微软雅黑"
FONT_EN = "Aptos"
TEAL = "0F766E"
LIGHT_TEAL = "E6FFFB"
PALE = "F8FAFC"
LINE = "D9E2EC"
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(100, 116, 139)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def style_paragraph(p, size: float = 10.5, bold: bool = False, color: RGBColor = INK, before: float = 0, after: float = 4, line: float = 1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    for run in p.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def add_para(doc: Document, text: str, size: float = 10.5, bold: bool = False, color: RGBColor = INK, before: float = 0, after: float = 4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=10.2, color=INK)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=15.5, bold=True, color=RGBColor(15, 118, 110))
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, bold=True, color=INK)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], LIGHT_TEAL)
        set_cell_border(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style_paragraph(p, size=9.8, bold=True, before=0, after=0)
    for r_index, row in enumerate(rows):
        cells = table.add_row().cells
        for c_index, value in enumerate(row):
            cells[c_index].text = value
            set_cell_border(cells[c_index])
            if r_index % 2 == 0:
                set_cell_shading(cells[c_index], PALE)
            cells[c_index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[c_index].paragraphs:
                style_paragraph(p, size=9.3, before=0, after=0)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_TEAL)
    set_cell_border(cell, "99F6E4")
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run_font(run, size=10.5, bold=True, color=RGBColor(15, 118, 110))
    p.paragraph_format.space_after = Pt(2)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(body)
    set_run_font(run2, size=10.2, color=INK)
    p2.paragraph_format.line_spacing = 1.12
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    styles = doc.styles
    styles["Normal"].font.name = FONT_EN
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = FONT_EN
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st.font.size = Pt(10.2)


def build_doc() -> Document:
    doc = Document()
    configure_doc(doc)

    # Cover-like compact title.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("食探（智健食探）商业计划书")
    set_run_font(run, size=22, bold=True, color=RGBColor(15, 23, 42))
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("AI 饮食记录与体重管理微信小程序 / 2026 年 4 月")
    set_run_font(run, size=10.5, color=MUTED)
    sub.paragraph_format.space_after = Pt(8)

    add_callout(
        doc,
        "一句话定位",
        "食探用 AI 降低饮食记录门槛，用趋势洞察和执行建议帮助用户长期管理体重。它不是一个单次“看热量”的小工具，而是围绕饮食、体重和行为调整形成闭环的轻量管理助手。",
    )

    add_table(
        doc,
        ["项目", "内容"],
        [
            ["品牌简称", "食探"],
            ["小程序名称建议", "智健食探"],
            ["核心关键词", "AI 饮食记录 / 体重管理 / 微信小程序 / 订阅制"],
            ["当前阶段", "主需求已获得早期验证，下一阶段聚焦订阅与留存验证"],
        ],
        [3.2, 13.2],
    )

    add_heading(doc, "一、为什么现在值得做", 1)
    add_para(doc, "中国用户的体重管理需求正在从“知道要减重”转向“需要有人帮我持续执行”。传统工具的核心问题不是数据不够，而是记录动作太重、反馈太慢、用户很难坚持。", after=3)
    add_bullets(doc, [
        "手动搜索食物、估算克重、填写份量，对普通用户来说门槛过高。",
        "很多 AI 识别产品停留在“拍照看热量”，没有把用户带入长期管理。",
        "体重变化需要周级、月级复盘，仅靠单日热量数字无法形成持续行动。",
        "微信小程序打开即用，适合饮食记录这种高频、低耐心场景。",
    ])
    add_callout(doc, "核心判断", "AI 识别会逐渐成为标配，真正的壁垒是用户是否愿意持续回来，并把自己的饮食和体重数据长期留在这里。")

    add_heading(doc, "二、用户痛点与解决方案", 1)
    add_table(
        doc,
        ["用户痛点", "食探方案"],
        [
            ["懒得手动记录", "拍照或一句话输入即可生成饮食记录，手动记录作为免费兜底。"],
            ["不知道今天吃得是否合理", "首页展示热量、三大营养素和餐次结构，让反馈变得即时。"],
            ["体重波动看不懂", "结合饮食、运动、体重趋势，形成周报复盘和调整建议。"],
            ["坚持几天就放弃", "用提醒、周报、目标进度和会员权益制造持续使用理由。"],
        ],
        [5.0, 11.4],
    )

    add_heading(doc, "三、产品结构", 1)
    add_para(doc, "产品主线围绕“记录 - 总览 - 复盘 - 调整”展开，先把最基础的饮食记录做轻，再逐步提高长期管理价值。", after=3)
    add_table(
        doc,
        ["模块", "功能", "用户价值"],
        [
            ["AI 饮食记录", "拍照识别、文字识别", "降低记录门槛，解决第一步很难开始的问题"],
            ["手动记录", "搜索食物、选择份量", "免费兜底，避免用户被 AI 次数或识别失败卡住"],
            ["首页总览", "热量、营养素、餐次、热力图", "让用户每天看到即时反馈"],
            ["身体数据", "体重、喝水、运动记录", "把饮食行为和身体变化连接起来"],
            ["周报复盘", "趋势洞察、问题归因、下周建议", "从工具升级为体重管理助手"],
            ["会员体系", "精准模式、完整历史、复盘报告", "承接持续价值与订阅收入"],
        ],
        [3.0, 5.1, 8.3],
    )

    add_heading(doc, "四、当前验证数据", 1)
    add_para(doc, "截至 2026 年 4 月，食探已经获得一批真实用户行为数据。当前阶段不把早期自愿付费样本作为主要卖点，重点看主链路是否有人使用、是否形成记录。", after=3)
    add_table(
        doc,
        ["指标", "当前数据", "说明"],
        [
            ["累计注册用户", "622", "已有早期用户池，可用于继续验证"],
            ["近 30 天活跃用户", "约 270", "说明不是纯一次性访问"],
            ["近 30 天 AI 分析任务", "3,408", "AI 记录入口有真实使用"],
            ["近 30 天饮食记录", "2,278", "用户会把分析结果落为记录"],
            ["分析后形成记录的用户转化率", "66.1%", "主链路转化表现较好"],
        ],
        [4.2, 3.0, 9.2],
    )
    add_callout(doc, "阶段结论", "当前数据说明“AI 降低记录门槛”这个需求成立。下一步要验证的是：用户是否会持续留下，并愿意为体重管理闭环订阅付费。")

    add_heading(doc, "五、商业模式：三档订阅 + 每日积分", 1)
    add_para(doc, "食探采用订阅会员制作为前台主心智，同时用“每日积分、当天清零”的方式控制 AI 成本。用户理解的是会员档位和每日可用额度，平台获得的是更稳定的订阅收入和更可控的单用户成本。", after=3)
    add_table(
        doc,
        ["会员档位", "月卡", "季卡", "年卡", "每日积分"],
        [
            ["轻度版", "9.9 元", "27.9 元", "99 元", "8 积分/天"],
            ["标准版", "19.9 元", "56.9 元", "199 元", "20 积分/天"],
            ["进阶版", "29.9 元", "84.9 元", "299 元", "40 积分/天"],
        ],
        [3.0, 3.0, 3.0, 3.0, 4.4],
    )
    add_para(doc, "页面可直接展示为：轻度版 9.9/月、27.9/季、99/年；标准版 19.9/月、56.9/季、199/年；进阶版 29.9/月、84.9/季、299/年。年卡和季卡页面展示“立省 xx 元”，引导长期订阅。", size=9.8, color=MUTED, after=3)
    add_heading(doc, "免费体验与积分规则", 2)
    add_table(
        doc,
        ["规则", "设计"],
        [
            ["新用户体验", "免费 3 天，每天 8 积分，当天清零，不累计"],
            ["积分有效期", "所有会员积分当天有效，不累计到次日"],
            ["运动记录", "1 积分 / 次"],
            ["基础记录 / 基础分析", "2 积分 / 次"],
            ["超额后", "等待次日恢复，或升级更高会员，或通过邀请好友获得额外积分"],
        ],
        [4.2, 12.2],
    )
    add_heading(doc, "增长奖励机制", 2)
    add_table(
        doc,
        ["机制", "规则", "控制方式"],
        [
            ["邀请奖励", "邀请人与被邀请人连续 3 天每天 +5 积分", "被邀请人完成 1 次有效使用后生效，每月设置上限防刷"],
            ["分享奖励", "生成分享海报奖励 1 积分", "建议每日上限 1 次，避免刷分和诱导分享风险"],
            ["续费机制", "支持自动续费，月卡/季卡/年卡同时提供", "页面突出长期订阅优惠和权益差异"],
        ],
        [3.2, 6.6, 6.6],
    )
    add_bullets(doc, [
        "前台卖订阅，不把用户教育成“每次都要算钱”的积分心智。",
        "积分当天清零，既能鼓励日常使用，又能避免低价年卡被重度用户长期透支。",
        "三档会员覆盖轻度记录、标准管理和高频进阶用户，价格梯度清晰。",
        "邀请和分享奖励用于低成本增长，但必须设置月度上限和有效使用门槛。",
    ])

    add_heading(doc, "六、竞争与差异化", 1)
    add_table(
        doc,
        ["竞品类型", "优势", "问题", "食探差异化"],
        [
            ["薄荷健康类", "数据库强、品牌成熟", "手动记录偏重，年轻用户坚持难", "AI 降低记录动作，把入口做轻"],
            ["Keep 类", "运动内容强、社区成熟", "饮食管理不是主线", "以饮食为核心，运动为补充"],
            ["AI 识别工具", "单次体验快", "缺少记录、趋势和管理闭环", "从识别进入长期管理"],
        ],
        [3.1, 3.7, 4.4, 5.2],
    )

    add_heading(doc, "七、90 天验证计划", 1)
    add_table(
        doc,
        ["方向", "动作", "目标"],
        [
            ["会员验证", "重做会员页，突出周报复盘、精准模式、趋势洞察", "验证用户是否愿意订阅"],
            ["留存提升", "上线周报复盘、目标进度、饮食归因建议", "让用户每周有回来理由"],
            ["内容获客", "小红书/抖音测试 10-20 条内容模板", "找到低成本自然获客方式"],
            ["数据看板", "补齐会员页曝光、点击、下单、支付结果埋点", "定位转化断点"],
        ],
        [3.1, 8.0, 5.3],
    )
    add_table(
        doc,
        ["关键指标", "目标"],
        [
            ["D30 留存", "提升至 20% 左右"],
            ["订阅付费率", "达到 3% 左右"],
            ["MAU", "突破 1,000"],
            ["内容获客", "跑出 1-2 个可复制内容模板"],
        ],
        [5.0, 11.4],
    )

    add_heading(doc, "八、风险与应对", 1)
    add_table(
        doc,
        ["风险", "表现", "应对"],
        [
            ["留存不足", "用户只体验 AI，不长期记录", "强化周报复盘、目标进度和提醒机制"],
            ["付费意愿不足", "用户认为只是普通工具", "权益从“次数”改为“体重管理结果”"],
            ["AI 成本不可控", "少数用户大量调用", "会员后台额度、模型分级和加量包兜底"],
            ["竞品跟进", "大平台上线类似识别功能", "聚焦本土饮食场景和更短闭环体验"],
        ],
        [3.5, 5.0, 7.9],
    )

    add_heading(doc, "九、结论", 1)
    add_para(doc, "食探（智健食探）当前不是要证明“AI 能不能识别食物”，而是要证明“AI 能否帮助用户长期坚持饮食和体重管理”。从早期数据看，主入口已经成立，分析到记录的转化也具备继续打磨的价值。", after=3)
    add_para(doc, "接下来 90 天，项目不应继续横向堆功能，而应围绕订阅会员制，把体重管理闭环做深：用 AI 记录降低门槛，用周报复盘建立价值，用数据看板验证留存和付费。", bold=True, color=RGBColor(15, 118, 110), after=8)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("食探（智健食探）商业计划书 · 2026 年 4 月")
    set_run_font(run, size=8.5, color=MUTED)

    return doc


def main() -> None:
    doc = build_doc()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
